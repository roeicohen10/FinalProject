from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer
import os, glob

# get final project path
DIR_PATH = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
EXPERIMENT_FOLDER_PATH = os.path.join(DIR_PATH, r'data\Experiments')

SINGLE_EXPERIMENT_PATH = os.path.join(DIR_PATH, r'data\exp_singel_report.docx')
DS_EXPERIMENT_PATH = os.path.join(DIR_PATH, r'data\dataset_exp_report.docx')
EMPTY_DOC_PATH = os.path.join(DIR_PATH, r'data\empty_doc.docx')
EMPTY_DOC__WITH_PARA_PATH = os.path.join(DIR_PATH, r'data\empty_doc_with_para.docx')

SINGLE_EXPERIMENT_FNAME = 'single_exp_report.docx'
DS_FNAME = 'ds_report.docx'

class Report:

    @classmethod
    def fill_se_table_rows(cls, table, **kwargs):
        rows = table.rows
        rows[1].cells[1].text = kwargs.get('ol_algo','')
        rows[2].cells[1].text = kwargs.get('ofs_algo','')
        rows[3].cells[1].text = kwargs.get('window_size','')
        rows[4].cells[1].text = kwargs.get('ol_runtime','')
        rows[5].cells[1].text = kwargs.get('ofs_runtime','')
        rows[6].cells[1].text = kwargs.get('accuracy','')
        rows[7].cells[1].text = kwargs.get('selected_features','')

        return table

    @classmethod
    def add_pic(cls, doc, image_path, width=4.52, height=3.4, paragraph_num=-1):
        paragraph = doc.paragraphs[paragraph_num]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width), height=Inches(height))
        run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return doc

    @classmethod
    def create_single_experiment_report(cls,**kwargs):
        doc = Document(SINGLE_EXPERIMENT_PATH)
        doc.tables[0] = cls.fill_se_table_rows(doc.tables[0], **kwargs)

        for image in [kwargs.get('first_image', ''),kwargs.get('second_image', ''),kwargs.get('third_image', ''),kwargs.get('forth_image', '')]:
            try:
                doc = cls.add_pic(doc, image)
            except FileNotFoundError as e:
                print(e)

        if kwargs.get('third_image', ''):
            for i in range(11):
                doc.add_paragraph('')
        doc.save(os.path.join(kwargs.get('export_path',''),SINGLE_EXPERIMENT_FNAME))

    @classmethod
    def fill_ds_ofs_runtime_table(cls, table,**kwargs):
        rows = table.rows
        rows[0].cells[1].text = kwargs.get('saola_runtime','-')
        rows[1].cells[1].text = kwargs.get('ai_runtime','-')
        rows[2].cells[1].text = kwargs.get('osfs_runtime', '-')
        rows[3].cells[1].text = kwargs.get('fosfs_runtime', '-')

    @classmethod
    def fill_ds_ol_runtime_table(cls, table,**kwargs):
        rows = table.rows
        rows[0].cells[1].text = kwargs.get('nn_runtime','-')
        rows[1].cells[1].text = kwargs.get('knn_3_runtime','-')
        rows[2].cells[1].text = kwargs.get('knn_5_runtime', '-')
        rows[3].cells[1].text = kwargs.get('nb_runtime', '-')
        rows[4].cells[1].text = kwargs.get('rf_runtime', '-')

        return table

    @classmethod
    def create_ds_report(cls, **kwargs):
        doc = Document(DS_EXPERIMENT_PATH)
        #set params for first 2 tables
        doc.tables[0].rows[0].cells[1].text =  kwargs.get('ds_name','')
        doc.tables[1].rows[2].cells[1].text = kwargs.get('window_sizes', '')

        paragraph_num = 6
        for image in [kwargs.get('wo_image', ''),kwargs.get('ai_image', ''),kwargs.get('osfs_image', ''),kwargs.get('fosfs_image', ''),kwargs.get('saola_image', '')]:
            try:
                doc = cls.add_pic(doc, image, width=4.58, height=2.8, paragraph_num=paragraph_num)
                paragraph_num += 2
                if paragraph_num == 10:
                    paragraph_num += 6

            except FileNotFoundError as e:
                print(e)
                paragraph_num += 1



        doc.tables[2] = cls.fill_ds_ofs_runtime_table(doc.tables[2], **kwargs)
        doc.tables[3] = cls.fill_ds_ol_runtime_table(doc.tables[3], **kwargs)
        doc.save(os.path.join(kwargs.get('export_path',''),DS_FNAME))

    @classmethod
    def combine_word_documents(cls,word_files, ds_name):
        merged_document = Document(EMPTY_DOC_PATH)
        composer = Composer(merged_document)
        for file in word_files:
            doc_temp = Document(file)
            composer.append(doc_temp)
        try:
            composer.save(os.path.join(os.path.dirname(word_files[0]),f'{ds_name}.docx'))
        except PermissionError as e:
            print(str(e))

    @classmethod
    def find_report_files(cls, ds_name, single_report=True):
        file_name = SINGLE_EXPERIMENT_FNAME if single_report else DS_FNAME

        return glob.glob(fr'{EXPERIMENT_FOLDER_PATH}\{ds_name}\**\{file_name}', recursive=True)


    @classmethod
    def combine_ds_experiments_reports(cls, ds_name):
        ds_report = cls.find_report_files(ds_name=ds_name, single_report=False)
        ds_report.append(EMPTY_DOC__WITH_PARA_PATH)
        single_reports = cls.find_report_files(ds_name=ds_name, single_report=True)
        ds_report.extend(single_reports)
        cls.combine_word_documents(ds_report, ds_name)

if __name__ == '__main__':
    params_one = {
        'ds_name':'FordA',
        'window_sizes': '100,200',
        'wo_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\FordA\Alpha Investing.png',
        'ai_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\FordA\Alpha Investing.png',
        'osfs_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\FordA\OSFS.png',
        'fosfs_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\FordA\Fast OSFS.png',
        'saola_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\FordA\SAOLA.png',
        'saola_runtime': str(520.2689666666658),
        'ai_runtime': str(1032.5714666666663),
        'osfs_runtime':str(1533.9393666666629),
        'fosfs_runtime':str(847.1284499999996),
        'nn_runtime': str(1.570100000001684),
        'knn_3_runtime': str(7.662487499998871),
        'knn_5_runtime': str(7.662487499998871),
        'nb_runtime': str(3.379575000000301),
        'rf_runtime': str(3.379575000000301),
        'export_path': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\FordA'
    }
    params_two = {
        'ol_algo': 'ANN (default)',
        'ofs_algo': 'AI (alpha=0.05, dw=0.05)',
        'window_size': str(100),
        'ol_runtime': '1.5020999999997287 ms',
        'ofs_runtime': '936.3674000000017 ms',
        'accuracy': str(0.487551867219917),
        'selected_features': '[0]',
        'first_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\ChlorineConcentration\100\SAOLA\K-Nearest Neighbors\acc_100_SAOLA_K-Nearest Neighbors.png',
        'second_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\ChlorineConcentration\100\SAOLA\K-Nearest Neighbors\acc_features_100_SAOLA_K-Nearest Neighbors.png',
        'third_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\ChlorineConcentration\100\SAOLA\K-Nearest Neighbors\acc_100_SAOLA_K-Nearest Neighbors.png',
        'forth_image': r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\ChlorineConcentration\100\SAOLA\K-Nearest Neighbors\acc_100_SAOLA_K-Nearest Neighbors.png',
        'export_path' : r'C:\Users\Roi\Documents\Degree\Semester 8\FinalProject\data\Experiments\ChlorineConcentration\100\SAOLA\Neural Network'
    }
    # Report.create_ds_report(**params_one)
    # Report.create_single_experiment_report(**params_two)
    Report.combine_ds_experiments_reports('ChlorineConcentration')
    # Report.find_report_files('ChlorineConcentration', single_report=False)


